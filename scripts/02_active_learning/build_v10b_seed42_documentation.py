from __future__ import annotations

import hashlib
import json
import math
import os
import shutil
import subprocess
import sys
import textwrap
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd


PROJECT_ROOT = Path(r"C:\Users\user\Desktop\vlm\Defect_VLM_Project")
DOC_DEPS = PROJECT_ROOT / ".codex_doc_deps"
if DOC_DEPS.exists():
    sys.path.insert(0, str(DOC_DEPS))

import matplotlib

matplotlib.use("Agg")
import matplotlib.font_manager as fm
import matplotlib.pyplot as plt


V10_RUN = PROJECT_ROOT / "runs" / "active_learning_ablation_v10_neu_large_pool" / "v10_neu_large_pool_smoke_20260712_185923"
V10B_PROBE = PROJECT_ROOT / "runs" / "v10b_selection_probe" / "v10b_selection_probe_20260712_194239"
V10B_SINGLE = PROJECT_ROOT / "runs" / "v10b_single_training" / "v10b_single_training_20260712_195514"
AUDIT = PROJECT_ROOT / "runs" / "v10b_independent_result_audit" / "v10b_independent_result_audit_20260712_202634"

RUNS_ROOT = PROJECT_ROOT / "runs" / "v10b_seed42_documentation"

ROUND0 = "__SHARED_ROUND0__"
RANDOM = "GTFreeRandom"
V9B = "DetectorInstanceRichDINOBalanced"
V10B = "DetectorUncertaintyDINOInstanceReducedV10b"

DISPLAY = {
    ROUND0: "Round0",
    RANDOM: "Random",
    V9B: "V9b",
    V10B: "V10b",
    "V9b": "V9b",
    "V10b": "V10b",
    "Random": "Random",
}

NEU6 = ["crazing", "inclusion", "patches", "pitted_surface", "rolled-in_scale", "scratches"]


def bool_env(name: str, default: bool = False) -> bool:
    raw = os.environ.get(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "y", "on"}


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def rows_count(path: Path) -> int | None:
    if path.suffix.lower() != ".csv":
        return None
    try:
        return int(len(pd.read_csv(path)))
    except Exception:
        return None


def source_registry() -> pd.DataFrame:
    files = [
        (V10_RUN / "config.json", "V10 smoke config"),
        (V10_RUN / "all_round_results.csv", "V10 aggregate metrics"),
        (V10_RUN / "actual_instance_statistics_by_round.csv", "V10 post-hoc XML statistics"),
        (V10_RUN / "actual_class_distribution_by_round.csv", "V10 cumulative class distribution"),
        (V10_RUN / "v9b_round1_detector_scores.csv", "V9b detector score saturation"),
        (V10_RUN / "development_eval_v10.csv", "development eval manifest"),
        (V10_RUN / "v10_neu_large_pool_smoke_summary.md", "V10 summary"),
        (V10B_PROBE / "config.json", "V10b probe config"),
        (V10B_PROBE / "v9b_selected_samples.csv", "V9b selected query"),
        (V10B_PROBE / "v10b_selected_samples.csv", "V10b selected query"),
        (V10B_PROBE / "v9b_v10b_selection_overlap.csv", "selection overlap"),
        (V10B_PROBE / "selection_component_summary.csv", "selection component summary"),
        (V10B_PROBE / "selection_geometry_summary.csv", "selection geometry summary"),
        (V10B_PROBE / "pool_score_saturation.csv", "pool score saturation"),
        (V10B_PROBE / "posthoc_actual_instance_stats.csv", "query batch XML statistics"),
        (V10B_PROBE / "posthoc_actual_class_distribution.csv", "query batch class distribution"),
        (V10B_PROBE / "v10b_selection_probe_summary.md", "probe summary"),
        (V10B_SINGLE / "config.json", "V10b single training config"),
        (V10B_SINGLE / "all_round_results.csv", "V10b aggregate metrics"),
        (V10B_SINGLE / "comparison_with_existing_v10_baselines.csv", "V10b baseline comparison"),
        (V10B_SINGLE / "actual_instance_statistics_by_round.csv", "V10b post-hoc XML statistics"),
        (V10B_SINGLE / "actual_class_distribution_by_round.csv", "V10b cumulative class distribution"),
        (V10B_SINGLE / "v10b_single_training_summary.md", "V10b summary"),
        (AUDIT / "implementation_integrity_audit.csv", "integrity audit"),
        (AUDIT / "recalculated_aggregate_metrics.csv", "recalculated aggregate metrics"),
        (AUDIT / "annotation_efficiency_recalculation.csv", "annotation efficiency"),
        (AUDIT / "cumulative_class_balance_metrics.csv", "class balance"),
        (AUDIT / "selection_change_recalculation.csv", "selection change"),
        (AUDIT / "recovered_per_class_metrics_v10.csv", "per-class recovery"),
        (AUDIT / "per_class_v10b_minus_random.csv", "per-class V10b-Random"),
        (AUDIT / "per_class_v10b_minus_v9b.csv", "per-class V10b-V9b"),
        (AUDIT / "per_class_recovery_registry.csv", "per-class recovery registry"),
        (AUDIT / "revised_evidence_classification.csv", "evidence classification"),
        (AUDIT / "revised_root_cause_evidence.csv", "root cause evidence"),
        (AUDIT / "recommended_next_step.md", "recommended next step"),
        (AUDIT / "v10_per_class_recovery_summary.md", "per-class summary"),
        (AUDIT / "v10b_independent_result_audit.md", "independent audit summary"),
    ]
    rows = []
    for path, purpose in files:
        rows.append(
            {
                "purpose": purpose,
                "path": str(path),
                "exists": path.exists(),
                "bytes": path.stat().st_size if path.exists() else 0,
                "sha256": sha256_file(path) if path.exists() else "",
                "rows": rows_count(path) if path.exists() else None,
            }
        )
    return pd.DataFrame(rows)


def assert_sources_present(registry: pd.DataFrame) -> None:
    missing = registry[~registry["exists"]]
    if len(missing):
        raise FileNotFoundError("Missing required source files:\n" + missing[["purpose", "path"]].to_string(index=False))


def norm_strategy(strategy: Any) -> str:
    return DISPLAY.get(str(strategy), str(strategy))


def load_metrics() -> dict[str, Any]:
    v10_cfg = read_json(V10_RUN / "config.json")
    probe_cfg = read_json(V10B_PROBE / "config.json")
    v10b_cfg = read_json(V10B_SINGLE / "config.json")
    audit_metrics = read_csv(AUDIT / "recalculated_aggregate_metrics.csv")
    v10_results = read_csv(V10_RUN / "all_round_results.csv")
    v10b_result = read_csv(V10B_SINGLE / "all_round_results.csv")
    aggregate = pd.concat([v10_results, v10b_result], ignore_index=True, sort=False)
    for col in ["map50", "map5095", "precision", "recall", "labeled_budget"]:
        aggregate[col] = pd.to_numeric(aggregate[col], errors="coerce")
    aggregate["strategy_display"] = aggregate["strategy"].map(norm_strategy)
    aggregate["f1"] = [
        2 * p * r / (p + r) if np.isfinite(p) and np.isfinite(r) and p + r else np.nan
        for p, r in zip(aggregate["precision"], aggregate["recall"])
    ]

    by_strategy = {row["strategy"]: row for _, row in aggregate.iterrows()}
    round0 = by_strategy[ROUND0]
    gains = []
    for strategy in [RANDOM, V9B, V10B]:
        row = by_strategy[strategy]
        gains.append(
            {
                "strategy": strategy,
                "strategy_display": norm_strategy(strategy),
                "map50_gain": row["map50"] - round0["map50"],
                "map5095_gain": row["map5095"] - round0["map5095"],
                "precision_gain": row["precision"] - round0["precision"],
                "recall_gain": row["recall"] - round0["recall"],
            }
        )
    gains_df = pd.DataFrame(gains)

    per_random = read_csv(AUDIT / "per_class_v10b_minus_random.csv")
    per_v9b = read_csv(AUDIT / "per_class_v10b_minus_v9b.csv")
    efficiency = read_csv(AUDIT / "annotation_efficiency_recalculation.csv")
    cumulative_balance = read_csv(AUDIT / "cumulative_class_balance_metrics.csv")
    selection_geometry = read_csv(V10B_PROBE / "selection_geometry_summary.csv")
    selection_components = read_csv(V10B_PROBE / "selection_component_summary.csv")
    overlap = read_csv(V10B_PROBE / "v9b_v10b_selection_overlap.csv")
    query_stats = read_csv(V10B_PROBE / "posthoc_actual_instance_stats.csv")
    v10_actual_stats = read_csv(V10_RUN / "actual_instance_statistics_by_round.csv")
    v10b_actual_stats = read_csv(V10B_SINGLE / "actual_instance_statistics_by_round.csv")
    actual_stats = pd.concat([v10_actual_stats, v10b_actual_stats], ignore_index=True, sort=False)
    v10_actual_class = read_csv(V10_RUN / "actual_class_distribution_by_round.csv")
    v10b_actual_class = read_csv(V10B_SINGLE / "actual_class_distribution_by_round.csv")
    actual_class = pd.concat([v10_actual_class, v10b_actual_class], ignore_index=True, sort=False)

    return {
        "v10_cfg": v10_cfg,
        "probe_cfg": probe_cfg,
        "v10b_cfg": v10b_cfg,
        "audit_metrics": audit_metrics,
        "aggregate": aggregate,
        "gains": gains_df,
        "per_random": per_random,
        "per_v9b": per_v9b,
        "efficiency": efficiency,
        "cumulative_balance": cumulative_balance,
        "selection_geometry": selection_geometry,
        "selection_components": selection_components,
        "overlap": overlap,
        "query_stats": query_stats,
        "actual_stats": actual_stats,
        "actual_class": actual_class,
    }


def validate_metrics(data: dict[str, Any]) -> pd.DataFrame:
    aggregate = data["aggregate"]
    audit = data["audit_metrics"]
    rows = []
    for strategy in [ROUND0, RANDOM, V9B, V10B]:
        src = aggregate[aggregate["strategy"].eq(strategy)].iloc[0]
        rec = audit[audit["strategy"].eq(strategy)]
        if rec.empty:
            rows.append({"metric_group": "aggregate", "strategy": strategy, "status": "review", "note": "not found in audit metrics"})
            continue
        rec = rec.iloc[0]
        for metric in ["map50", "map5095", "precision", "recall"]:
            diff = abs(float(src[metric]) - float(rec[metric]))
            rows.append(
                {
                    "metric_group": "aggregate",
                    "strategy": strategy,
                    "metric": metric,
                    "source_value": float(src[metric]),
                    "audit_value": float(rec[metric]),
                    "abs_diff": diff,
                    "status": "pass" if diff < 1e-6 else "review",
                }
            )
    per_random = data["per_random"]
    rows.append(
        {
            "metric_group": "per_class",
            "strategy": "V10b-Random",
            "metric": "neu6_count",
            "source_value": len(per_random),
            "audit_value": 6,
            "abs_diff": abs(len(per_random) - 6),
            "status": "pass" if len(per_random) == 6 else "review",
        }
    )
    return pd.DataFrame(rows)


def setup_matplotlib() -> None:
    font_names = [f.name for f in fm.fontManager.ttflist]
    for candidate in ["Malgun Gothic", "맑은 고딕", "NanumGothic", "Arial Unicode MS", "DejaVu Sans"]:
        if candidate in font_names:
            plt.rcParams["font.family"] = candidate
            break
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 150
    plt.rcParams["savefig.dpi"] = 220


def save_fig(fig: plt.Figure, path: Path, title: str, registry: list[dict[str, Any]]) -> None:
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    registry.append({"figure": path.name, "title": title, "path": str(path), "status": "generated"})


def value_label(ax: plt.Axes, fmt: str = "{:.3f}") -> None:
    for patch in ax.patches:
        height = patch.get_height()
        if not np.isfinite(height):
            continue
        ax.annotate(
            fmt.format(height),
            (patch.get_x() + patch.get_width() / 2, height),
            ha="center",
            va="bottom" if height >= 0 else "top",
            fontsize=8,
            xytext=(0, 2 if height >= 0 else -8),
            textcoords="offset points",
        )


def plot_grouped_bars(ax: plt.Axes, labels: list[str], series: dict[str, list[float]], ylabel: str) -> None:
    x = np.arange(len(labels))
    width = 0.8 / max(1, len(series))
    for i, (name, vals) in enumerate(series.items()):
        ax.bar(x - 0.4 + width / 2 + i * width, vals, width, label=name)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=20, ha="right")
    ax.set_ylabel(ylabel)
    ax.legend(fontsize=8, ncol=2)
    ax.grid(axis="y", alpha=0.25)


def minmax_norm(values: list[float]) -> list[float]:
    arr = np.asarray(values, dtype=float)
    if not len(arr) or np.nanmax(arr) - np.nanmin(arr) <= 1e-12:
        return [0.0 for _ in values]
    return ((arr - np.nanmin(arr)) / (np.nanmax(arr) - np.nanmin(arr))).tolist()


def generate_figures(data: dict[str, Any], fig_dir: Path) -> pd.DataFrame:
    setup_matplotlib()
    fig_dir.mkdir(parents=True, exist_ok=True)
    registry: list[dict[str, Any]] = []

    aggregate = data["aggregate"].copy()
    aggregate = aggregate.set_index("strategy")
    order = [ROUND0, RANDOM, V9B, V10B]
    labels = [norm_strategy(s) for s in order]
    metrics = ["map50", "map5095", "precision", "recall"]

    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    plot_grouped_bars(ax, labels, {m: [float(aggregate.loc[s, m]) for s in order] for m in metrics}, "score")
    ax.set_title("Figure 1. Aggregate detector performance")
    save_fig(fig, fig_dir / "fig01_aggregate_performance.png", "Aggregate detector performance grouped bar chart", registry)

    gains = data["gains"].set_index("strategy")
    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    gain_metrics = ["map50_gain", "map5095_gain", "precision_gain", "recall_gain"]
    plot_grouped_bars(ax, [norm_strategy(s) for s in [RANDOM, V9B, V10B]], {m: [float(gains.loc[s, m]) for s in [RANDOM, V9B, V10B]] for m in gain_metrics}, "gain vs Round0")
    ax.axhline(0, color="#333333", linewidth=0.8)
    ax.set_title("Figure 2. Gain over shared Round0")
    save_fig(fig, fig_dir / "fig02_gain_vs_round0.png", "Round0 대비 gain", registry)

    fig, ax = plt.subplots(figsize=(6.4, 4.8))
    for s in order:
        row = aggregate.loc[s]
        ax.scatter(float(row["recall"]), float(row["precision"]), s=85)
        ax.annotate(norm_strategy(s), (float(row["recall"]), float(row["precision"])), xytext=(6, 4), textcoords="offset points")
    ax.set_xlabel("Recall")
    ax.set_ylabel("Precision")
    ax.set_title("Figure 3. Precision-Recall scatter")
    ax.grid(alpha=0.25)
    save_fig(fig, fig_dir / "fig03_precision_recall.png", "Precision-Recall scatter", registry)

    geom = data["selection_geometry"].copy()
    comp = data["selection_components"].copy()
    rows = []
    for strategy in ["Random", "V9b", "V10b"]:
        g = geom[geom["strategy"].eq(strategy)].iloc[0]
        c = comp[comp["strategy"].eq(strategy)].set_index("component")
        rows.append(
            {
                "strategy": strategy,
                "pairwise cosine": float(g["batch_pairwise_cosine_similarity_mean"]),
                "distance to initial": float(g["distance_to_initial_mean"]),
                "detector uncertainty": float(c.loc["detector_uncertainty", "mean"]),
                "pseudo box count": float(c.loc["detector_pseudo_box_count", "mean"]),
            }
        )
    geom_df = pd.DataFrame(rows).set_index("strategy")
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    normed = pd.DataFrame({col: minmax_norm(geom_df[col].tolist()) for col in geom_df.columns}, index=geom_df.index)
    plot_grouped_bars(ax, normed.index.tolist(), {col: normed[col].tolist() for col in normed.columns}, "min-max normalized")
    ax.set_title("Figure 4. Selection geometry and component shifts (normalized)")
    save_fig(fig, fig_dir / "fig04_selection_geometry.png", "Selection geometry normalized plot", registry)

    overlap = data["overlap"].iloc[0]
    fig, ax = plt.subplots(figsize=(6.4, 3.8))
    categories = ["overlap", "V9b only", "V10b only"]
    values = [float(overlap["overlap_count"]), float(overlap["v9b_only_count"]), float(overlap["v10b_only_count"])]
    ax.bar(categories, values, color=["#6B7280", "#4F81BD", "#70AD47"])
    ax.set_ylabel("images")
    ax.set_title(f"Figure 5. V9b/V10b query overlap (Jaccard={float(overlap['jaccard']):.3f})")
    value_label(ax, "{:.0f}")
    save_fig(fig, fig_dir / "fig05_selection_overlap.png", "V9b/V10b overlap", registry)

    query = data["query_stats"].copy()
    query["strategy"] = query["source_strategy"].map(norm_strategy)
    qmetrics = {
        "total bbox": query["total_bbox_instances"].astype(float).tolist(),
        "bbox / image": query["mean_bbox_per_image"].astype(float).tolist(),
        "multi-object ratio": query["multi_object_image_ratio"].astype(float).tolist(),
    }
    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    normalized = {k: minmax_norm(v) for k, v in qmetrics.items()}
    plot_grouped_bars(ax, query["strategy"].tolist(), normalized, "min-max normalized")
    ax.set_title("Figure 6. Query batch actual bbox statistics (normalized)")
    save_fig(fig, fig_dir / "fig06_query_instance_statistics.png", "Query batch actual bbox statistics", registry)

    actual_class = data["actual_class"].copy()
    actual_class["strategy_display"] = actual_class["strategy"].map(norm_strategy)
    actual_class["actual_xml_class"] = actual_class["actual_xml_class"].map(lambda x: str(x).replace("rolled-in-scale", "rolled-in_scale"))
    dist = actual_class[actual_class["strategy"].isin([RANDOM, V9B, V10B])]
    pivot = dist.pivot_table(index="actual_xml_class", columns="strategy_display", values="bbox_instance_count", aggfunc="sum").reindex(NEU6)
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    plot_grouped_bars(ax, NEU6, {col: pivot[col].fillna(0).astype(float).tolist() for col in ["Random", "V9b", "V10b"] if col in pivot.columns}, "bbox instances")
    ax.set_title("Figure 7. Cumulative actual class distribution")
    save_fig(fig, fig_dir / "fig07_cumulative_class_distribution.png", "Cumulative actual class distribution", registry)

    per_random = data["per_random"].copy()
    fig, ax = plt.subplots(figsize=(8.6, 4.6))
    vals = per_random["V10b_minus_Random_ap5095"].astype(float)
    colors = ["#70AD47" if v >= 0 else "#C0504D" for v in vals]
    ax.bar(per_random["class_name"], vals, color=colors)
    ax.axhline(0, color="#333333", linewidth=0.8)
    ax.set_ylabel("AP50-95 difference")
    ax.set_title("Figure 8. Per-class V10b minus Random AP50-95")
    ax.tick_params(axis="x", rotation=25)
    value_label(ax, "{:.3f}")
    save_fig(fig, fig_dir / "fig08_per_class_v10b_minus_random.png", "Per-class V10b minus Random AP50-95", registry)

    per_v9b = data["per_v9b"].copy()
    fig, ax = plt.subplots(figsize=(8.6, 4.6))
    vals = per_v9b["V10b_minus_V9b_ap5095"].astype(float)
    ax.bar(per_v9b["class_name"], vals, color="#70AD47")
    ax.axhline(0, color="#333333", linewidth=0.8)
    ax.set_ylabel("AP50-95 difference")
    ax.set_title("Figure 9. Per-class V10b minus V9b AP50-95")
    ax.tick_params(axis="x", rotation=25)
    value_label(ax, "{:.3f}")
    save_fig(fig, fig_dir / "fig09_per_class_v10b_minus_v9b.png", "Per-class V10b minus V9b AP50-95", registry)

    eff = data["efficiency"].copy()
    eff = eff[eff["strategy"].isin([RANDOM, V9B, V10B])].copy()
    eff["strategy_display"] = eff["strategy"].map(norm_strategy)
    raw_metrics = {
        "additional bbox": eff["additional_bbox_vs_round0"].astype(float).tolist(),
        "mAP50-95 gain": eff["map5095_gain_vs_round0"].astype(float).tolist(),
        "gain / bbox": eff["map5095_gain_per_additional_bbox"].astype(float).tolist(),
    }
    fig, ax = plt.subplots(figsize=(8.4, 4.5))
    plot_grouped_bars(ax, eff["strategy_display"].tolist(), {k: minmax_norm(v) for k, v in raw_metrics.items()}, "min-max normalized")
    ax.set_title("Figure 10. Annotation efficiency indicators (normalized)")
    save_fig(fig, fig_dir / "fig10_annotation_efficiency.png", "Annotation efficiency", registry)

    fig, ax = plt.subplots(figsize=(7.8, 3.8))
    ax.axis("off")
    boxes = [
        ("Directly verified facts", 0.12, "#D9EAF7"),
        ("Supported interpretations", 0.46, "#E2F0D9"),
        ("Unverified hypotheses", 0.80, "#FFF2CC"),
    ]
    for text, x, color in boxes:
        rect = plt.Rectangle((x, 0.32), 0.24, 0.28, facecolor=color, edgecolor="#555555")
        ax.add_patch(rect)
        ax.text(x + 0.12, 0.46, text, ha="center", va="center", fontsize=10, wrap=True)
    ax.annotate("", xy=(0.46, 0.46), xytext=(0.36, 0.46), arrowprops={"arrowstyle": "->"})
    ax.annotate("", xy=(0.80, 0.46), xytext=(0.70, 0.46), arrowprops={"arrowstyle": "->"})
    ax.set_title("Figure 11. Evidence level separation")
    save_fig(fig, fig_dir / "fig11_evidence_levels.png", "Evidence level diagram", registry)

    fig, ax = plt.subplots(figsize=(9.2, 3.8))
    ax.axis("off")
    steps = [
        "V9b failure",
        "V10b weight revision",
        "Selection-only gate",
        "Seed42 detector gate",
        "Per-class audit",
        "Seed43-46 validation\nrunning / pending",
    ]
    xs = np.linspace(0.05, 0.85, len(steps))
    for i, (x, step) in enumerate(zip(xs, steps)):
        rect = plt.Rectangle((x, 0.36), 0.12, 0.28, facecolor="#E8EEF5", edgecolor="#4F81BD")
        ax.add_patch(rect)
        ax.text(x + 0.06, 0.50, step, ha="center", va="center", fontsize=8, wrap=True)
        if i < len(steps) - 1:
            ax.annotate("", xy=(xs[i + 1], 0.50), xytext=(x + 0.12, 0.50), arrowprops={"arrowstyle": "->", "color": "#333333"})
    ax.set_title("Figure 12. Experiment evolution flow")
    save_fig(fig, fig_dir / "fig12_experiment_evolution.png", "Experiment evolution flow", registry)

    return pd.DataFrame(registry)


def fmt(x: Any, ndigits: int = 6) -> str:
    try:
        if not np.isfinite(float(x)):
            return ""
        return f"{float(x):.{ndigits}f}"
    except Exception:
        return str(x)


def markdown_table(df: pd.DataFrame, columns: list[str] | None = None, ndigits: int = 6) -> str:
    sub = df[columns].copy() if columns else df.copy()
    for col in sub.columns:
        if pd.api.types.is_numeric_dtype(sub[col]):
            sub[col] = sub[col].map(lambda x: fmt(x, ndigits))
    return sub.to_markdown(index=False)


def build_markdown(data: dict[str, Any], fig_registry: pd.DataFrame, out_dir: Path) -> str:
    aggregate = data["aggregate"].copy()
    aggregate_table = aggregate[["strategy_display", "labeled_budget", "map50", "map5095", "precision", "recall", "f1"]].rename(
        columns={"strategy_display": "Strategy", "labeled_budget": "Budget", "map50": "mAP50", "map5095": "mAP50-95", "precision": "Precision", "recall": "Recall", "f1": "F1"}
    )
    per_random = data["per_random"][["class_name", "V10b_minus_Random_ap5095"]].copy()
    per_v9b = data["per_v9b"][["class_name", "V10b_minus_V9b_ap5095"]].copy()
    efficiency = data["efficiency"].copy()
    efficiency["strategy_display"] = efficiency["strategy"].map(norm_strategy)
    eff_table = efficiency[efficiency["strategy"].isin([RANDOM, V9B, V10B])][
        ["strategy_display", "total_bbox", "additional_bbox_vs_round0", "map5095_gain_vs_round0", "map5095_gain_per_additional_bbox"]
    ].rename(
        columns={
            "strategy_display": "Strategy",
            "total_bbox": "Total bbox",
            "additional_bbox_vs_round0": "Additional bbox",
            "map5095_gain_vs_round0": "mAP50-95 gain",
            "map5095_gain_per_additional_bbox": "Gain/additional bbox",
        }
    )

    v10b = aggregate[aggregate["strategy"].eq(V10B)].iloc[0]
    random = aggregate[aggregate["strategy"].eq(RANDOM)].iloc[0]
    v9b = aggregate[aggregate["strategy"].eq(V9B)].iloc[0]

    lines = [
        "# V10b: Detector Uncertainty와 Visual Diversity를 강화한 산업 결함 탐지 Active Learning 샘플 선택 전략",
        "",
        "**부제:** NEU-DET Large-Pool One-Cycle 실험 및 Random Baseline 비교 - Seed 42 Development Gate 검증 보고서",
        "",
        "- 연구 주제: object detection active learning에서 annotation 효율을 높이는 acquisition 전략 검증",
        "- 대상 데이터셋: NEU-DET 전체 1,800장",
        "- detector: YOLOv8n",
        "- 평가 단계: development gate",
        "- Final test status: locked / unused",
        f"- Experiment IDs: `{data['v10_cfg']['experiment_id']}`, `{data['probe_cfg'].get('status', 'selection_probe')}`, `{data['v10b_cfg']['experiment_id']}`, audit `v10b_independent_result_audit_20260712_202634`",
        "",
        "## 1. Executive Summary",
        "",
        f"V10b는 acquisition seed 42, training seed 1042의 development one-cycle에서 Random보다 높은 aggregate 성능을 보였다. mAP50-95는 V10b `{fmt(v10b['map5095'])}`, Random `{fmt(random['map5095'])}`로 차이는 `{fmt(v10b['map5095'] - random['map5095'])}`이다. Precision과 recall도 각각 `{fmt(v10b['precision'] - random['precision'])}`, `{fmt(v10b['recall'] - random['recall'])}` 높았다.",
        "",
        "다만 seed42는 방법 개발과 가중치 조정에 사용된 development seed이므로, 이 결과는 일반화 주장이나 final evidence가 아니다. 현재 증거 수준은 exploratory success / development gate pass이며, 다음 단계는 V10b weight를 고정한 seed43-46 Random 대비 one-cycle 검증이다.",
        "",
        "## 2. 실험 프로토콜",
        "",
        "- Dataset: NEU-DET 전체 1,800장",
        "- Acquisition pool: 900장, 클래스당 150장",
        "- Development eval: 300장, 클래스당 50장",
        "- Final test: 300장, locked / unused",
        "- Unused reserve: 300장",
        "- Initial labeled set: 60장",
        "- Query: 30장",
        "- Final labeled budget: 90장",
        "- Acquisition seed: 42",
        "- Training seed: 1042",
        "- Detector: YOLOv8n, 100 epochs, imgsz 640",
        "- Pool/dev/final overlap: 0",
        "",
        "Acquisition score는 다음 형태로 정리된다.",
        "",
        "$$S(x)=w_u U(x)+w_d D(x)+w_b B(x)+w_i I(x)$$",
        "",
        "- `U(x)`: detector uncertainty",
        "- `D(x)`: DINO visual distance/diversity",
        "- `B(x)`: predicted class deficit",
        "- `I(x)`: pseudo instance count score",
        "",
        "각 component는 iterative selection 과정에서 정규화되어 사용된다. 특히 pseudo-instance score는 상한에서 포화될 수 있으므로, 단순히 bbox가 많은 샘플을 더 고르는 것이 detector utility를 보장하지 않는다.",
        "",
        "## 3. V9b 실패 분석",
        "",
        "V9b는 Random보다 instance-rich하고 class entropy가 높았지만, detector 성능은 낮았다. V9b의 precision은 높았으나 recall이 낮아 보수적인 detector가 되었고, 이는 높은 pseudo-instance weight가 detector utility와 직접 연결되지 않았음을 시사한다.",
        "",
        "핵심 차이는 '많은 객체 수'와 '학습 가능한 유효 객체 수'의 차이다. V9b는 bbox instance를 더 많이 확보했지만 mAP50-95 gain per additional bbox가 가장 낮았다.",
        "",
        "## 4. V10b 설계 수정",
        "",
        "| Component | V9b | V10b | Direction |",
        "|---|---:|---:|---|",
        "| detector uncertainty | 0.20 | 0.25 | increase |",
        "| DINO visual distance | 0.25 | 0.35 | increase |",
        "| predicted class deficit | 0.15 | 0.15 | keep |",
        "| pseudo instance count | 0.40 | 0.25 | decrease |",
        "",
        "V10b의 설계 가정은 instance-rich 편향을 낮추고, initial set과 중복되지 않는 시각적 coverage 및 detector decision boundary가 불확실한 샘플을 더 확보하는 것이다. 이 해석은 현재 supported interpretation이지 causal proof는 아니다.",
        "",
        "## 5. Selection-only probe 결과",
        "",
        "V10b는 V9b와 30개 중 18개가 겹쳤고, 12개가 교체되었다. Jaccard는 0.428571로, 실제로 다른 selector로 동작했음을 보여준다.",
        "",
        "![Figure 4](figures/fig04_selection_geometry.png)",
        "",
        "![Figure 5](figures/fig05_selection_overlap.png)",
        "",
        "## 6. Detector 성능 결과",
        "",
        markdown_table(aggregate_table, ndigits=6),
        "",
        "![Figure 1](figures/fig01_aggregate_performance.png)",
        "",
        "![Figure 2](figures/fig02_gain_vs_round0.png)",
        "",
        "![Figure 3](figures/fig03_precision_recall.png)",
        "",
        f"V10b minus Random: mAP50 `{fmt(v10b['map50'] - random['map50'])}`, mAP50-95 `{fmt(v10b['map5095'] - random['map5095'])}`, precision `{fmt(v10b['precision'] - random['precision'])}`, recall `{fmt(v10b['recall'] - random['recall'])}`.",
        "",
        f"V10b minus V9b: mAP50 `{fmt(v10b['map50'] - v9b['map50'])}`, mAP50-95 `{fmt(v10b['map5095'] - v9b['map5095'])}`, precision `{fmt(v10b['precision'] - v9b['precision'])}`, recall `{fmt(v10b['recall'] - v9b['recall'])}`.",
        "",
        "## 7. Per-class AP 분석",
        "",
        "V10b는 Random 대비 AP50-95에서 crazing, patches, pitted_surface, scratches에서 이겼고 inclusion, rolled-in_scale에서는 졌다.",
        "",
        markdown_table(per_random, ndigits=4),
        "",
        "![Figure 8](figures/fig08_per_class_v10b_minus_random.png)",
        "",
        "V10b는 V9b 대비 NEU 6개 클래스 모두에서 AP50-95가 높았다.",
        "",
        markdown_table(per_v9b, ndigits=4),
        "",
        "![Figure 9](figures/fig09_per_class_v10b_minus_v9b.png)",
        "",
        "## 8. Annotation Efficiency",
        "",
        markdown_table(eff_table, ndigits=6),
        "",
        "![Figure 10](figures/fig10_annotation_efficiency.png)",
        "",
        "Gain per additional bbox는 annotation time을 직접 측정한 비용 모델이 아니다. 동일 이미지 내 bbox 수가 annotation cost에 미치는 영향을 단순 proxy로 본 보조 지표이며, 단독 성능지표로 사용해서는 안 된다.",
        "",
        "## 9. Cumulative labeled set 분석",
        "",
        "V10b cumulative 90장의 실제 XML class distribution은 crazing 30, inclusion 34, patches 48, pitted_surface 30, rolled-in_scale 34, scratches 30이며 total bbox는 206, entropy는 2.5620이다. Query batch만 보면 scratches 편향과 patches 부족이 보이지만 initial set과 결합된 cumulative set은 비교적 균형적이다.",
        "",
        "![Figure 6](figures/fig06_query_instance_statistics.png)",
        "",
        "![Figure 7](figures/fig07_cumulative_class_distribution.png)",
        "",
        "높은 entropy가 성능 향상의 직접 원인이라는 주장은 금지한다. 현재는 class balance가 detector utility와 양립했다는 정도만 말할 수 있다.",
        "",
        "## 10. Integrity Audit",
        "",
        "독립 감사 결과는 pass with caveat이다. initial size 60, query size 30, cumulative 90, train/dev overlap 0, final test used=False/read=False, V10b training count=1, selection SHA match, acquisition seed 42, training seed 1042, aggregate metric recovery 일치가 확인되었다.",
        "",
        "주의점은 git_dirty=True, seed42가 tuning/development seed, single training seed, YAML nc=15 placeholder 문제이다. 본 문서의 per-class 해석은 NEU 6개 클래스만 사용한다.",
        "",
        "## 11. 증거 수준 구분",
        "",
        "![Figure 11](figures/fig11_evidence_levels.png)",
        "",
        "### Directly verified facts",
        "",
        "- V10b seed42 mAP50-95=0.340866",
        "- Random seed42 mAP50-95=0.329700",
        "- validation-only recovery가 aggregate 기록과 일치",
        "- final test 미사용",
        "- selection Jaccard=0.428571",
        "",
        "### Supported interpretations",
        "",
        "- V10b가 V9b의 instance-rich 편향을 완화했다.",
        "- V10b selection은 visual redundancy를 줄였다.",
        "- V10b는 candidate method로서 seed43-46 검증을 수행할 가치가 있다.",
        "",
        "### Unverified hypotheses",
        "",
        "- uncertainty 증가가 recall 향상의 직접 원인이다.",
        "- DINO diversity 증가가 성능 향상의 직접 원인이다.",
        "- class entropy 증가가 성능 향상의 직접 원인이다.",
        "- V10b가 모든 seed나 final test에서도 Random보다 우수하다.",
        "",
        "## 12. 최종 판정과 다음 단계",
        "",
        "V10b는 V9b의 instance-rich 편향을 완화하고 uncertainty와 visual diversity를 강화함으로써 recall 저하를 회복했으며, seed42 development one-cycle에서 Random보다 높은 mAP, precision, recall과 더 나은 annotation-efficiency proxy를 달성했다. 다만 seed42는 방법 선택에 사용된 development seed이므로, V10b 가중치를 고정한 뒤 새로운 acquisition seed에서 재현성을 검증해야 한다.",
        "",
        "![Figure 12](figures/fig12_experiment_evolution.png)",
        "",
        "다음 단계는 V10b weights frozen 상태에서 seed43-46 Random vs V10b one-cycle 검증이며, 그 결과가 통과한 뒤 full learning curve와 final test 1회 평가로 진행한다.",
        "",
        "> 본 보고서는 acquisition seed 42와 training seed 1042를 사용한 development-stage 결과를 정리한 것이며 독립 multiseed 검증과 locked final-test 평가 이전의 탐색/방법개발 단계 증거이다.",
    ]
    return "\n".join(lines) + "\n"


def write_tables(data: dict[str, Any], out_dir: Path) -> pd.DataFrame:
    out_dir.mkdir(parents=True, exist_ok=True)
    tables = []
    table_map = {
        "table01_aggregate_detector_performance.csv": data["aggregate"],
        "table02_gain_vs_round0.csv": data["gains"],
        "table03_per_class_v10b_minus_random.csv": data["per_random"],
        "table04_per_class_v10b_minus_v9b.csv": data["per_v9b"],
        "table05_annotation_efficiency.csv": data["efficiency"],
        "table06_cumulative_class_balance.csv": data["cumulative_balance"],
        "table07_integrity_audit.csv": read_csv(AUDIT / "implementation_integrity_audit.csv"),
    }
    for name, df in table_map.items():
        path = out_dir / name
        df.to_csv(path, index=False, encoding="utf-8-sig")
        tables.append({"table": name, "path": str(path), "rows": len(df), "status": "generated"})
    return pd.DataFrame(tables)


def import_docx_modules():
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
        return {
            "Document": Document,
            "WD_SECTION": WD_SECTION,
            "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
            "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "OxmlElement": OxmlElement,
            "qn": qn,
            "Inches": Inches,
            "Pt": Pt,
            "RGBColor": RGBColor,
        }
    except Exception as exc:
        raise RuntimeError(
            "python-docx is unavailable. Install it with: .\\.python311\\python.exe -m pip install --target .codex_doc_deps python-docx"
        ) from exc


def set_run_font(run, font_name: str, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_docx_table(doc, df: pd.DataFrame, caption: str, max_rows: int | None = None) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cap = doc.add_paragraph(caption)
    cap.style = "Caption"
    rows = df.head(max_rows).copy() if max_rows else df.copy()
    table = doc.add_table(rows=1, cols=len(rows.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(rows.columns):
        hdr[i].text = str(col)
        shade_cell(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, "Malgun Gothic", 8.5, bold=True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for _, row in rows.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(rows.columns):
            val = row[col]
            if isinstance(val, (float, np.floating)):
                text = fmt(val, 4)
            else:
                text = str(val)
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, "Malgun Gothic", 8)
    doc.add_paragraph()


def build_docx(data: dict[str, Any], md_text: str, fig_registry: pd.DataFrame, table_registry: pd.DataFrame, out_dir: Path) -> Path:
    mods = import_docx_modules()
    Document = mods["Document"]
    WD_ALIGN_PARAGRAPH = mods["WD_ALIGN_PARAGRAPH"]
    Inches = mods["Inches"]
    Pt = mods["Pt"]
    RGBColor = mods["RGBColor"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(mods["qn"]("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string("1F4D78")

    header = section.header.paragraphs[0]
    header.text = "V10b Seed42 Development Gate Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        set_run_font(r, "Malgun Gothic", 9, "666666")
    footer = section.footer.paragraphs[0]
    footer.text = "Page "
    add_page_number(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        set_run_font(r, "Malgun Gothic", 9, "666666")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("V10b: Detector Uncertainty와 Visual Diversity를 강화한\n산업 결함 탐지 Active Learning 샘플 선택 전략")
    set_run_font(tr, "Malgun Gothic", 22, "0B2545", bold=True)
    title.paragraph_format.space_after = Pt(8)
    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("Seed 42 Development Gate 검증 보고서 | NEU-DET Large-Pool One-Cycle 실험 및 Random Baseline 비교")
    set_run_font(sr, "Malgun Gothic", 11.5, "555555")
    subtitle.paragraph_format.space_after = Pt(14)

    meta = pd.DataFrame(
        [
            ("Dataset", "NEU-DET 1,800 images"),
            ("Protocol", "Pool 900 / Dev 300 / Final test 300 locked"),
            ("Seed", "acquisition 42 / training 1042"),
            ("Detector", "YOLOv8n, 100 epochs, imgsz 640"),
            ("Evidence level", "Development gate; not final evidence"),
            ("Final test", "Locked / unused"),
        ],
        columns=["Field", "Value"],
    )
    add_docx_table(doc, meta, "Table 1. Report metadata")

    doc.add_heading("목차", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. 실험 프로토콜",
        "3. V9b 실패 분석",
        "4. V10b 설계 수정",
        "5. Selection-only probe 결과",
        "6. Detector 성능 결과",
        "7. Per-class AP 분석",
        "8. Annotation Efficiency",
        "9. Cumulative labeled set 분석",
        "10. Integrity Audit",
        "11. 증거 수준 구분",
        "12. 최종 판정과 다음 단계",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_page_break()

    sections = [
        ("1. Executive Summary", [
            "V10b는 seed42 development one-cycle에서 Random보다 높은 aggregate mAP, precision, recall을 보였다. 그러나 seed42는 방법 개발에 사용된 tuning/development seed이므로 일반화 증거는 아니다.",
            "V10b는 V9b 대비 NEU 6개 클래스 모두에서 AP50-95가 높았고, Random 대비로는 4개 클래스에서 이기고 2개 클래스에서 졌다.",
            "다음 단계는 V10b weight를 고정한 seed43-46 Random 대비 one-cycle 검증이다.",
        ]),
        ("2. 실험 프로토콜", [
            "NEU-DET 전체 1,800장을 사용했고, acquisition pool 900장, development eval 300장, final test 300장 locked/unused, unused reserve 300장으로 구성했다.",
            "Acquisition score: S(x)=w_uU(x)+w_dD(x)+w_bB(x)+w_iI(x). U는 detector uncertainty, D는 DINO visual distance, B는 predicted class deficit, I는 pseudo instance score이다.",
        ]),
        ("3. V9b 실패 분석", [
            "V9b는 bbox instance와 class entropy를 높였지만 Random보다 낮은 mAP와 recall을 보였다. 이는 instance-rich 선택이 항상 detector utility로 이어지지 않음을 보여준다.",
        ]),
        ("4. V10b 설계 수정", [
            "V10b는 pseudo instance weight를 0.40에서 0.25로 낮추고, DINO visual distance를 0.25에서 0.35로, detector uncertainty를 0.20에서 0.25로 높였다.",
            "이 설계는 과도한 instance-rich 편향을 완화하고 시각적 coverage와 불확실 샘플 확보를 강화하려는 후보 전략이다.",
        ]),
    ]
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    doc.add_heading("5. Selection-only probe 결과", level=1)
    doc.add_paragraph("V10b는 V9b와 30개 중 18개가 겹쳤고, 12개가 교체되었다. Jaccard는 0.428571이다.")
    for fig_name in ["fig04_selection_geometry.png", "fig05_selection_overlap.png"]:
        doc.add_picture(str(out_dir / "figures" / fig_name), width=Inches(6.1))
        doc.add_paragraph(fig_name.replace(".png", ""), style="Caption")

    doc.add_heading("6. Detector 성능 결과", level=1)
    agg = data["aggregate"][["strategy_display", "labeled_budget", "map50", "map5095", "precision", "recall", "f1"]].rename(
        columns={"strategy_display": "Strategy", "labeled_budget": "Budget"}
    )
    add_docx_table(doc, agg, "Table 2. V10 Large-Pool Seed42 Development 성능 비교")
    for fig_name in ["fig01_aggregate_performance.png", "fig02_gain_vs_round0.png", "fig03_precision_recall.png"]:
        doc.add_picture(str(out_dir / "figures" / fig_name), width=Inches(6.1))
        doc.add_paragraph(fig_name.replace(".png", ""), style="Caption")

    doc.add_heading("7. Per-class AP 분석", level=1)
    doc.add_paragraph("V10b는 Random 대비 AP50-95에서 crazing, patches, pitted_surface, scratches에서 이겼고 inclusion, rolled-in_scale에서는 졌다.")
    add_docx_table(doc, data["per_random"][["class_name", "V10b_minus_Random_ap5095"]], "Table 3. Per-class V10b minus Random AP50-95")
    doc.add_picture(str(out_dir / "figures" / "fig08_per_class_v10b_minus_random.png"), width=Inches(6.1))
    doc.add_paragraph("Figure 8. Per-class V10b minus Random AP50-95", style="Caption")
    add_docx_table(doc, data["per_v9b"][["class_name", "V10b_minus_V9b_ap5095"]], "Table 4. Per-class V10b minus V9b AP50-95")
    doc.add_picture(str(out_dir / "figures" / "fig09_per_class_v10b_minus_v9b.png"), width=Inches(6.1))
    doc.add_paragraph("Figure 9. Per-class V10b minus V9b AP50-95", style="Caption")

    doc.add_heading("8. Annotation Efficiency", level=1)
    eff = data["efficiency"].copy()
    eff["strategy_display"] = eff["strategy"].map(norm_strategy)
    add_docx_table(
        doc,
        eff[eff["strategy"].isin([RANDOM, V9B, V10B])][
            ["strategy_display", "total_bbox", "additional_bbox_vs_round0", "map5095_gain_vs_round0", "map5095_gain_per_additional_bbox"]
        ],
        "Table 5. Annotation efficiency proxy",
    )
    doc.add_picture(str(out_dir / "figures" / "fig10_annotation_efficiency.png"), width=Inches(6.1))
    doc.add_paragraph("Figure 10. Annotation efficiency indicators", style="Caption")
    doc.add_paragraph("Gain per additional bbox는 실제 annotation time을 측정한 비용 모델이 아니며 보조 proxy로만 해석해야 한다.")

    doc.add_heading("9. Cumulative labeled set 분석", level=1)
    doc.add_paragraph("V10b cumulative 90장은 total bbox 206, entropy 2.5620으로 기록되었다. Query batch 편향과 cumulative set 균형은 구분해서 해석해야 한다.")
    for fig_name in ["fig06_query_instance_statistics.png", "fig07_cumulative_class_distribution.png"]:
        doc.add_picture(str(out_dir / "figures" / fig_name), width=Inches(6.1))
        doc.add_paragraph(fig_name.replace(".png", ""), style="Caption")

    doc.add_heading("10. Integrity Audit", level=1)
    doc.add_paragraph("독립 감사 판정은 pass with caveat이다. final test는 사용되지 않았고, validation-only recovery는 aggregate 기록과 일치했다.")
    audit = read_csv(AUDIT / "implementation_integrity_audit.csv")
    add_docx_table(doc, audit[["check", "status", "evidence"]].head(14), "Table 6. Implementation integrity audit", max_rows=14)

    doc.add_heading("11. 증거 수준 구분", level=1)
    doc.add_picture(str(out_dir / "figures" / "fig11_evidence_levels.png"), width=Inches(6.1))
    doc.add_paragraph("Figure 11. Evidence level separation", style="Caption")
    for title, items in [
        ("Directly verified facts", ["V10b seed42 mAP50-95=0.340866", "Random seed42 mAP50-95=0.329700", "final test unused"]),
        ("Supported interpretations", ["V10b가 V9b instance-rich 편향을 완화", "visual redundancy 감소", "candidate method 가치"]),
        ("Unverified hypotheses", ["모든 seed에서 Random보다 우수", "final test에서도 우수", "DINO diversity가 직접 원인"]),
    ]:
        doc.add_heading(title, level=2)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

    doc.add_heading("12. 최종 판정과 다음 단계", level=1)
    doc.add_paragraph(
        "V10b는 seed42 development gate를 통과했지만, 일반화 증거는 아니다. V10b weight를 고정하고 seed43-46 one-cycle 검증을 수행한 뒤 full learning curve 및 final test 1회 평가로 진행해야 한다."
    )
    doc.add_picture(str(out_dir / "figures" / "fig12_experiment_evolution.png"), width=Inches(6.1))
    doc.add_paragraph("Figure 12. Experiment evolution flow", style="Caption")
    final_note = doc.add_paragraph()
    run = final_note.add_run(
        "본 보고서는 acquisition seed 42와 training seed 1042를 사용한 development-stage 결과를 정리한 것이며 독립 multiseed 검증과 locked final-test 평가 이전의 탐색/방법개발 단계 증거이다."
    )
    set_run_font(run, "Malgun Gothic", 10.5, "9B1C1C", bold=True)

    docx_path = out_dir / "V10b_Seed42_Development_Gate_Updated_Report.docx"
    doc.save(docx_path)
    return docx_path


def try_pdf_export(docx_path: Path, out_dir: Path) -> tuple[Path | None, str]:
    pdf_path = out_dir / "V10b_Seed42_Development_Gate_Updated_Report.pdf"
    candidates = ["soffice", "libreoffice"]
    for exe in candidates:
        if shutil.which(exe):
            try:
                subprocess.run(
                    [exe, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
                    cwd=str(out_dir),
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=120,
                )
                produced = out_dir / f"{docx_path.stem}.pdf"
                if produced.exists() and produced.stat().st_size > 0:
                    if produced != pdf_path:
                        produced.replace(pdf_path)
                    return pdf_path, "success"
            except Exception as exc:
                return None, f"{exe} failed: {exc!r}"
    return None, "LibreOffice/soffice not found"


def main() -> None:
    t0 = time.perf_counter()
    dry_run = bool_env("AL_DOC_DRY_RUN_ONLY", True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = RUNS_ROOT / f"v10b_seed42_documentation_{timestamp}"
    fig_dir = out_dir / "figures"
    tables_dir = out_dir / "tables"
    out_dir.mkdir(parents=True, exist_ok=True)

    registry = source_registry()
    assert_sources_present(registry)
    registry.to_csv(out_dir / "report_source_metrics.csv", index=False, encoding="utf-8-sig")

    final_test_path = V10_RUN / "final_test_v10_LOCKED_UNUSED.csv"
    final_test_exists = final_test_path.exists()

    data = load_metrics()
    validation = validate_metrics(data)
    validation.to_csv(out_dir / "report_recalculated_metrics.csv", index=False, encoding="utf-8-sig")

    config = {
        "experiment_id": out_dir.name,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "project_root": str(PROJECT_ROOT),
        "dry_run": dry_run,
        "source_runs": {
            "v10_smoke": str(V10_RUN),
            "v10b_selection_probe": str(V10B_PROBE),
            "v10b_single_training": str(V10B_SINGLE),
            "v10b_independent_audit": str(AUDIT),
        },
        "final_test_file_exists": final_test_exists,
        "final_test_read": False,
        "multiseed_running_outputs_read": False,
        "design_preset": "standard_business_brief",
        "header_template": "memo_masthead",
        "doc_deps": str(DOC_DEPS),
        "source_files": registry.to_dict(orient="records"),
    }
    (out_dir / "report_generation_config.json").write_text(json.dumps(config, ensure_ascii=False, indent=2, default=str), encoding="utf-8")

    planned_figures = [
        "fig01_aggregate_performance.png",
        "fig02_gain_vs_round0.png",
        "fig03_precision_recall.png",
        "fig04_selection_geometry.png",
        "fig05_selection_overlap.png",
        "fig06_query_instance_statistics.png",
        "fig07_cumulative_class_distribution.png",
        "fig08_per_class_v10b_minus_random.png",
        "fig09_per_class_v10b_minus_v9b.png",
        "fig10_annotation_efficiency.png",
        "fig11_evidence_levels.png",
        "fig12_experiment_evolution.png",
    ]
    if dry_run:
        pd.DataFrame({"figure": planned_figures, "status": "planned"}).to_csv(out_dir / "report_figure_registry.csv", index=False, encoding="utf-8-sig")
        pd.DataFrame({"table": [f"table{i:02d}" for i in range(1, 8)], "status": "planned"}).to_csv(out_dir / "report_table_registry.csv", index=False, encoding="utf-8-sig")
        print("=" * 100)
        print("[DRY RUN] V10b seed42 documentation preflight complete")
        print(f"Output dir: {out_dir}")
        print("source runs verified=True")
        print("final test read=False")
        print("multiseed running outputs read=False")
        print(f"planned figures count={len(planned_figures)}")
        print("DOCX generated=False")
        print("=" * 100)
        return

    fig_registry = generate_figures(data, fig_dir)
    fig_registry.to_csv(out_dir / "report_figure_registry.csv", index=False, encoding="utf-8-sig")
    table_registry = write_tables(data, tables_dir)
    table_registry.to_csv(out_dir / "report_table_registry.csv", index=False, encoding="utf-8-sig")
    md_text = build_markdown(data, fig_registry, out_dir)
    md_path = out_dir / "V10b_Seed42_Development_Gate_Updated_Report.md"
    md_path.write_text(md_text, encoding="utf-8")
    docx_path = build_docx(data, md_text, fig_registry, table_registry, out_dir)
    pdf_path, pdf_status = try_pdf_export(docx_path, out_dir)

    config.update(
        {
            "generated_figures_count": int(len(fig_registry)),
            "generated_tables_count": int(len(table_registry)),
            "markdown_path": str(md_path),
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path) if pdf_path else None,
            "pdf_status": pdf_status,
            "runtime_sec": time.perf_counter() - t0,
        }
    )
    (out_dir / "report_generation_config.json").write_text(json.dumps(config, ensure_ascii=False, indent=2, default=str), encoding="utf-8")

    aggregate_status = "pass" if validation[validation["metric_group"].eq("aggregate")]["status"].eq("pass").all() else "review"
    per_class_status = "pass" if validation[validation["metric_group"].eq("per_class")]["status"].eq("pass").all() else "review"

    print("=" * 100)
    print("[DONE] V10b seed42 documentation generated")
    print("source runs verified=True")
    print("final test read=False")
    print("multiseed running outputs read=False")
    print(f"generated figures count={len(fig_registry)}")
    print(f"generated tables count={len(table_registry)}")
    print(f"DOCX path={docx_path}")
    print(f"PDF path={pdf_path if pdf_path else 'unavailable'}")
    print(f"PDF unavailable reason={'' if pdf_path else pdf_status}")
    print(f"aggregate metric validation status={aggregate_status}")
    print(f"per-class metric validation status={per_class_status}")
    print("seed42 evidence level=development gate / exploratory success")
    print("next experiment=seed43-46 frozen V10b validation")
    print(f"Output dir={out_dir}")
    print("=" * 100)


if __name__ == "__main__":
    main()
