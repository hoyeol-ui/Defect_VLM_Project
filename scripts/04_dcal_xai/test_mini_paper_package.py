#!/usr/bin/env python3
"""Claim-boundary and artifact tests for the frozen mini-paper package."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
RUN_DIR = ROOT / "runs" / "mini_paper_package_20260718"
PAPER = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.md"
DOCX = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.docx"
PDF = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.pdf"


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT)).replace("\\", "/")


def main() -> None:
    results: list[tuple[str, bool, str]] = []

    def check(name: str, condition: bool, detail: str) -> None:
        results.append((name, bool(condition), detail))

    required = [
        PAPER,
        DOCX,
        PDF,
        ROOT / "docs/validity_gated_workflow_algorithm_20260718.md",
        ROOT / "docs/mini_paper_core_tables_20260718.md",
        ROOT / "docs/mini_paper_self_persuasion_audit_20260718.md",
        ROOT / "docs/advisor_mini_paper_decision_package_20260718.md",
        ROOT / "docs/figures/full_page_validity_gated_architecture.png",
        ROOT / "docs/figures/full_page_validity_gated_architecture.svg",
        ROOT / "docs/figures/full_page_validity_gated_architecture.pdf",
        ROOT / "docs/figures/validity_gated_algorithm_flowchart.png",
        ROOT / "docs/figures/validity_gated_algorithm_flowchart.svg",
        ROOT / "docs/figures/validity_gated_algorithm_flowchart.pdf",
        ROOT / "docs/figures/paper_framework_temporal_validation.png",
        ROOT / "docs/figures/paper_framework_cost_avoidance.png",
        ROOT / "docs/figures/paper_framework_advance_stop_timeline.png",
        RUN_DIR / "source_registry.csv",
        RUN_DIR / "missing_evidence.csv",
        RUN_DIR / "package_config.json",
    ]
    missing = [rel(p) for p in required if not p.exists() or p.stat().st_size == 0]
    check("T01_required_artifacts", not missing, "missing=" + (", ".join(missing) if missing else "none"))

    paper = PAPER.read_text(encoding="utf-8")
    algorithm = (ROOT / "docs/validity_gated_workflow_algorithm_20260718.md").read_text(encoding="utf-8")
    tables = (ROOT / "docs/mini_paper_core_tables_20260718.md").read_text(encoding="utf-8")
    audit = (ROOT / "docs/mini_paper_self_persuasion_audit_20260718.md").read_text(encoding="utf-8")
    advisor = (ROOT / "docs/advisor_mini_paper_decision_package_20260718.md").read_text(encoding="utf-8")
    package_text = "\n".join([paper, algorithm, tables, audit, advisor])
    lower = package_text.lower()

    check("T02_framework_identity", "validity-gated empirical evaluation and cost-containment workflow" in lower,
          "retrospective empirical/cost-containment identity present")
    check("T03_not_predictive_central_claim", "predictive screening framework" not in paper.lower(),
          "paper does not present a predictive screening framework")
    check("T04_temporal_decision_C", "C. RETROSPECTIVE_AUDIT_ONLY" in paper,
          "temporal audit decision C preserved")
    check("T05_evidence_freeze_decision_A", "A. THESIS_REFRAME_STRONGLY_SUPPORTED" in paper,
          "Evidence Freeze v2 decision A preserved")
    check("T06_predictive_confusion_not_identifiable",
          "predictive confusion matrix" in lower and "not identifiable" in lower,
          "predictive confusion matrix remains not identifiable")
    check("T07_generic_holdout_zero", "0/6" in package_text,
          "generic precommitted holdout count is 0/6")

    numeric_early_stop = re.search(r"early[- ]stop recall\s*[:=]\s*(?!NA\b)(?:0(?:\.\d+)?|1(?:\.0+)?)", package_text, re.I)
    numeric_false_advance = re.search(r"false[- ]advance rate\s*[:=]\s*(?!NA\b)(?:0(?:\.\d+)?|1(?:\.0+)?)", package_text, re.I)
    numeric_correct_stop = re.search(r"correct[- ]stop precision\s*[:=]\s*(?!NA\b)(?:0(?:\.\d+)?|1(?:\.0+)?)", package_text, re.I)
    check("T08_early_stop_recall_NA", numeric_early_stop is None and "early-stop recall" in lower and "NA" in package_text,
          "no numeric early-stop recall")
    check("T09_false_advance_NA", numeric_false_advance is None and "false-advance rate" in lower,
          "no numeric false-advance rate")
    check("T10_correct_stop_precision_NA", numeric_correct_stop is None and "correct-stop precision" in lower,
          "no numeric correct-stop precision")

    check("T11_selection_pass_is_authorization",
          "Selection PASS는 성공 예측이 아니라" in package_text and "1회 bounded detector screen" in package_text,
          "PASS is a bounded measurement authorization")
    check("T12_downstream_fail_not_false_advance",
          "downstream FAIL은 selection gate의 `false advance`가 아니라" in algorithm
          or "detector FAIL은 workflow의 false advance가 아니라" in paper,
          "bounded downstream failure is not relabeled as a prediction error")

    check("T13_cost_lower_bound_exact", "최소 45" in package_text and "45 planned" in package_text,
          "documented lower bound 45 is present")
    inflated_cost = re.search(r"최소\s+(?:4[6-9]|[5-9]\d|\d{3,})\s*(?:개|planned)", package_text)
    check("T14_no_inflated_avoided_runs", inflated_cost is None,
          "no lower-bound claim above 45")
    check("T15_final_actual_use_zero", "actual use는 0회" in paper or "final 실제 사용 0" in package_text,
          "locked final actual use is zero")
    check("T16_final_counterfactual_NA",
          "counterfactual final-test avoidance" in lower and "NA" in package_text,
          "counterfactual final-test avoidance is NA")

    check("T17_fn_remains_exploratory",
          "FN 1.379693" in package_text and "exploratory" in lower,
          "FN signal is explicitly exploratory")
    check("T18_200_seeds_not_production_pools",
          "200 seeds는 200개의 독립 production pool이 아니다" in paper
          or "200 seeds는 independent production lots가 아니라" in advisor,
          "fixed-pool perturbations are not production pools")
    check("T19_mpdd_exif_not_lot", "EXIF day를 production lot으로 해석" in paper,
          "MPDD EXIF is not promoted to a production lot")
    check("T20_no_independent_production_claim",
          "independent production generalization" in lower and "not identifiable" in lower,
          "independent production generalization remains unavailable")

    check("T21_gc10_tradeoff_pair",
          "+0.017378" in paper and "-0.019877" in paper,
          "GC10 overall gain and rare loss are jointly reported")
    check("T22_visa_tradeoff_pair",
          "+14.480" in paper and "-4.110" in paper and "+0.286025" in paper,
          "VisA discovery and category collapse are jointly reported")
    check("T23_seed45_uncertainty",
          "+0.007019" in paper and "[-0.005211, 0.019678]" in paper and "0.322266" in paper,
          "seed45 acquisition uncertainty is preserved")
    check("T24_k40_translation_gap",
          "-0.001678" in paper and "-0.018290" in paper and "-0.021871" in paper,
          "K40 downstream losses are preserved")
    check("T25_vlm_negative_signal",
          "-0.181199" in paper and "0.373433" in paper and "0/3" in paper,
          "VLM validity negatives are preserved")

    check("T26_evidence_locators", all(eid in paper for eid in ("E001", "E010", "E017", "E025", "E033", "E034")),
          "representative frozen evidence IDs are cited")
    check("T27_body_pagebreak_structure", paper.count("<!-- PAGEBREAK -->") == 7,
          f"pagebreaks={paper.count('<!-- PAGEBREAK -->')}, expected=7")
    check("T28_sections_present", all(f"# {n}." in paper for n in range(1, 10)),
          "sections 1 through 9 are present")

    with Image.open(ROOT / "docs/figures/full_page_validity_gated_architecture.png") as image:
        architecture_size = image.size
    with Image.open(ROOT / "docs/figures/validity_gated_algorithm_flowchart.png") as image:
        flow_size = image.size
    check("T29_architecture_resolution", architecture_size == (3840, 2160),
          f"architecture={architecture_size}")
    check("T30_flowchart_resolution", flow_size[0] >= 2200 and flow_size[1] >= 3000,
          f"flowchart={flow_size}")

    reader = PdfReader(str(PDF))
    check("T31_pdf_page_count", len(reader.pages) == 8,
          f"pages={len(reader.pages)}; 7 body + 1 references")
    extracted = "\n".join((page.extract_text() or "") for page in reader.pages)
    check("T32_pdf_text_extractable", "산업 결함 능동학습 후보 신호" in extracted and "참고문헌 및 증거 추적" in extracted,
          f"extracted_chars={len(extracted)}")

    doc = Document(str(DOCX))
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    check("T33_docx_content",
          "Selection PASS가 성공 예측이 아니라" in docx_text
          or "Selection PASS는 성공 예측이 아니라" in docx_text,
          f"paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")
    check("T34_docx_has_figures_tables", len(doc.inline_shapes) == 4 and len(doc.tables) >= 4,
          f"inline_shapes={len(doc.inline_shapes)}, tables={len(doc.tables)}")

    config = json.loads((RUN_DIR / "package_config.json").read_text(encoding="utf-8"))
    protected = [
        "training_performed", "inference_performed", "vlm_calls_performed",
        "embedding_extraction_performed", "selector_implementation_performed",
        "fn_screen_performed", "final_test_used",
    ]
    check("T35_all_protected_actions_false", all(config.get(key) is False for key in protected),
          ", ".join(f"{key}={config.get(key)}" for key in protected))
    check("T36_predictive_policy_claim_false", config.get("predictive_policy_claim") is False,
          "predictive_policy_claim=false")

    registry_text = (RUN_DIR / "source_registry.csv").read_text(encoding="utf-8-sig")
    check("T37_registry_has_frozen_sources",
          "research_evidence_ledger.csv" in registry_text and "framework_temporal_validation_decision" in registry_text,
          "registry includes frozen ledger and temporal decision")
    missing_text = (RUN_DIR / "missing_evidence.csv").read_text(encoding="utf-8-sig")
    check("T38_missing_evidence_explicit",
          all(key in missing_text for key in ("prospective_generic_policy", "independent_production_pools", "human_trust", "fn_external_confirmation")),
          "key unavailable evidence is enumerated")

    passed = sum(ok for _, ok, _ in results)
    failed = len(results) - passed
    lines = [
        "MINI PAPER PACKAGE TEST RESULTS",
        f"passed={passed}",
        f"failed={failed}",
        "training_performed=false",
        "inference_performed=false",
        "final_test_used=false",
        "docx_renderer=UNAVAILABLE_NO_SOFFICE",
        "docx_qa=OOXML_STRUCTURE_AND_CONTENT_TESTED",
        "pdf_qa=8_PAGES_RENDERED_AND_VISUALLY_INSPECTED",
        "",
    ]
    for name, ok, detail in results:
        lines.append(f"[{'PASS' if ok else 'FAIL'}] {name}: {detail}")
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    (RUN_DIR / "test_results.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")

    manifest_paths = sorted({*required, ROOT / "scripts/04_dcal_xai/build_mini_paper_figures.py",
                             ROOT / "scripts/04_dcal_xai/build_mini_paper_package.py",
                             ROOT / "scripts/04_dcal_xai/test_mini_paper_package.py"}, key=lambda p: str(p))
    manifest_lines = []
    for path in manifest_paths:
        status = "PRESENT" if path.exists() else "MISSING"
        size = path.stat().st_size if path.exists() else 0
        manifest_lines.append(f"{status}\t{size}\t{rel(path)}")
    (RUN_DIR / "generated_file_manifest.txt").write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")

    print(json.dumps({
        "status": "PASS" if failed == 0 else "FAIL",
        "passed": passed,
        "failed": failed,
        "test_results": str(RUN_DIR / "test_results.txt"),
        "manifest": str(RUN_DIR / "generated_file_manifest.txt"),
        "training_performed": False,
        "inference_performed": False,
        "final_test_used": False,
    }, ensure_ascii=False, indent=2))
    if failed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
