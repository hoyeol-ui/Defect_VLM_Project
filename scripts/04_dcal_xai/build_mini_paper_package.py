#!/usr/bin/env python3
"""Build the frozen-evidence mini-paper DOCX and its audit registry.

This script performs document generation only.  It does not train a model,
run inference, extract embeddings, implement a selector, screen false
negatives, or touch a locked final split.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph as RLParagraph,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.md"
DEFAULT_OUTPUT = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.docx"
DEFAULT_PDF = ROOT / "docs" / "mini_paper_validity_gated_industrial_al_20260718.pdf"
DEFAULT_RUN_DIR = ROOT / "runs" / "mini_paper_package_20260718"

NAVY = "173A5E"
BLUE = "2E6F9E"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF4EC"
PALE_AMBER = "FBF3E5"
PALE_RED = "FAECEC"
MID_GRAY = "697586"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=65, start=80, bottom=65, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="Malgun Gothic", size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=8, color=MID_GRAY)


def set_paragraph_border(paragraph, color=BLUE, size=12, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(8.6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(2.3)
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 16.5, NAVY, 0, 5),
        ("Heading 1", 13.2, NAVY, 6, 4),
        ("Heading 2", 10.6, BLUE, 4, 2.5),
        ("Heading 3", 9.4, BLUE, 3, 2),
    ):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Paper Subtitle" not in styles:
        styles.add_style("Paper Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle = styles["Paper Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(9.5)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor.from_string(MID_GRAY)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(3)

    if "Paper Meta" not in styles:
        styles.add_style("Paper Meta", WD_STYLE_TYPE.PARAGRAPH)
    meta = styles["Paper Meta"]
    meta.font.name = "Malgun Gothic"
    meta._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    meta.font.size = Pt(8)
    meta.font.color.rgb = RGBColor.from_string(MID_GRAY)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)

    if "Figure Caption" not in styles:
        styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Figure Caption"]
    caption.font.name = "Malgun Gothic"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    caption.font.size = Pt(7.4)
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(3)

    if "Compact List" not in styles:
        styles.add_style("Compact List", WD_STYLE_TYPE.PARAGRAPH)
    compact = styles["Compact List"]
    compact.base_style = normal
    compact.paragraph_format.left_indent = Cm(0.45)
    compact.paragraph_format.first_line_indent = Cm(-0.22)
    compact.paragraph_format.space_after = Pt(1.3)

    if "Reference" not in styles:
        styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    ref = styles["Reference"]
    ref.base_style = normal
    ref.font.size = Pt(8)
    ref.paragraph_format.left_indent = Cm(0.55)
    ref.paragraph_format.first_line_indent = Cm(-0.55)
    ref.paragraph_format.space_after = Pt(1.4)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline(paragraph, text: str, *, base_size=None) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=NAVY)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Cascadia Mono", size=(base_size or 8.2) - 0.3, color="7A3E00")
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def image_width_for(path: Path) -> float:
    name = path.name
    return {
        "full_page_validity_gated_architecture.png": 7.0,
        "full_page_validity_gated_architecture_v2.png": 7.0,
        "deeppcb_prospective_stop_case.png": 6.5,
        "validity_gated_algorithm_flowchart.png": 3.55,
        "discovery_composition_utility_matrix.png": 6.5,
        "paper_framework_temporal_validation.png": 6.2,
        "paper_framework_cost_avoidance.png": 6.25,
    }.get(name, 6.2)


def add_image(doc: Document, path: Path, alt: str) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        add_inline(p, f"[그림 누락: {alt} — {path}]", base_size=8)
        set_paragraph_shading(p, PALE_RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(path), width=Inches(image_width_for(path)))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    usable_cm = 18.0
    weights = [1.0] * ncols
    if ncols == 2:
        weights = [1.15, 4.85]
    elif ncols == 4:
        weights = [1.15, 1.75, 2.05, 1.85]
    elif ncols == 5:
        weights = [1.15, 1.65, 1.85, 1.75, 1.55]
    total = sum(weights)
    widths = [usable_cm * w / total for w in weights]
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx in range(ncols):
            cell = row.cells[c_idx]
            cell.width = Cm(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            value = values[c_idx] if c_idx < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, base_size=6.9 if ncols >= 5 else 7.2)
            for run in p.runs:
                if r_idx == 0:
                    set_run_font(run, size=7.0, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_border(p, BLUE)
    set_paragraph_shading(p, PALE_BLUE)
    add_inline(p, text.strip(), base_size=8.3)


def build_docx(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.40)
    section.left_margin = Cm(1.50)
    section.right_margin = Cm(1.50)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "VALIDITY-GATED INDUSTRIAL DEFECT ACTIVE LEARNING"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, name="Aptos", size=7.2, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    doc.core_properties.title = "산업 결함 능동학습 후보 신호의 단계적 타당성 평가"
    doc.core_properties.subject = "Validity-gated empirical evaluation and cost containment"
    doc.core_properties.author = "Defect VLM Project"
    doc.core_properties.keywords = "industrial defect, active learning, validity gate, cost containment"

    i = 0
    seen_title = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            add_image(doc, image_path, image_match.group(1))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith(">"):
            add_blockquote(doc, line.lstrip("> "))
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not seen_title:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(p, text, base_size=16.5)
                seen_title = True
            else:
                p = doc.add_heading(level=level)
                add_inline(p, text)
            i += 1
            continue
        if seen_title and line.startswith("**Validity-Gated Evaluation"):
            p = doc.add_paragraph(style="Paper Subtitle")
            add_inline(p, line.strip("*"), base_size=9.5)
            i += 1
            continue
        if seen_title and line.startswith("연구 방향 승인용"):
            p = doc.add_paragraph(style="Paper Meta")
            add_inline(p, line, base_size=8)
            i += 1
            continue
        if line.startswith("**그림 "):
            p = doc.add_paragraph(style="Figure Caption")
            add_inline(p, line, base_size=7.4)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="Compact List")
            p.paragraph_format.left_indent = Cm(0.52)
            p.paragraph_format.first_line_indent = Cm(-0.38)
            add_inline(p, line, base_size=8.3)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="Compact List")
            add_inline(p, "• " + line[2:], base_size=8.2)
            i += 1
            continue
        style = "Reference" if re.match(r"^\[\d+\]", line) else "Normal"
        p = doc.add_paragraph(style=style)
        add_inline(p, line)
        i += 1

    # Compatibility and print settings.
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    compat_setting = OxmlElement("w:compatSetting")
    compat_setting.set(qn("w:name"), "compatibilityMode")
    compat_setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    compat_setting.set(qn("w:val"), "15")
    compat.append(compat_setting)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def markdown_inline_to_reportlab(text: str) -> str:
    """Escape text and retain the small inline subset used by the paper."""
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`(.+?)`", r'<font name="MalgunGothic" color="#7A3E00">\1</font>', escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", escaped)
    return escaped


def register_pdf_fonts() -> None:
    regular = Path(r"C:\Windows\Fonts\malgun.ttf")
    bold = Path(r"C:\Windows\Fonts\malgunbd.ttf")
    if not regular.exists() or not bold.exists():
        raise FileNotFoundError("Malgun Gothic font files are required for Korean PDF output")
    if "MalgunGothic" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("MalgunGothic", str(regular)))
        pdfmetrics.registerFont(TTFont("MalgunGothic-Bold", str(bold)))
        pdfmetrics.registerFontFamily(
            "MalgunGothic",
            normal="MalgunGothic",
            bold="MalgunGothic-Bold",
            italic="MalgunGothic",
            boldItalic="MalgunGothic-Bold",
        )


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "KoreanBody", parent=base["BodyText"], fontName="MalgunGothic",
            fontSize=7.35, leading=9.15, alignment=TA_JUSTIFY,
            spaceAfter=2.2, wordWrap="CJK", textColor=colors.HexColor("#263646"),
        ),
        "title": ParagraphStyle(
            "KoreanTitle", fontName="MalgunGothic-Bold", fontSize=15.2,
            leading=19, alignment=TA_CENTER, spaceAfter=4, wordWrap="CJK",
            textColor=colors.HexColor("#173A5E"),
        ),
        "subtitle": ParagraphStyle(
            "EnglishSubtitle", fontName="MalgunGothic", fontSize=8.2,
            leading=10, alignment=TA_CENTER, spaceAfter=3,
            textColor=colors.HexColor("#697586"),
        ),
        "meta": ParagraphStyle(
            "PaperMeta", fontName="MalgunGothic", fontSize=7.0,
            leading=8.5, alignment=TA_CENTER, spaceAfter=4,
            textColor=colors.HexColor("#697586"),
        ),
        "h1": ParagraphStyle(
            "KoreanH1", fontName="MalgunGothic-Bold", fontSize=11.5,
            leading=14, spaceBefore=3.5, spaceAfter=3, keepWithNext=True,
            wordWrap="CJK", textColor=colors.HexColor("#173A5E"),
        ),
        "h2": ParagraphStyle(
            "KoreanH2", fontName="MalgunGothic-Bold", fontSize=9.2,
            leading=11.5, spaceBefore=2.5, spaceAfter=2, keepWithNext=True,
            wordWrap="CJK", textColor=colors.HexColor("#2E6F9E"),
        ),
        "h3": ParagraphStyle(
            "KoreanH3", fontName="MalgunGothic-Bold", fontSize=8.3,
            leading=10.5, spaceBefore=2.5, spaceAfter=1.5, keepWithNext=True,
            wordWrap="CJK", textColor=colors.HexColor("#2E6F9E"),
        ),
        "quote": ParagraphStyle(
            "KoreanQuote", fontName="MalgunGothic", fontSize=7.3,
            leading=9.2, alignment=TA_JUSTIFY, leftIndent=10, rightIndent=8,
            borderColor=colors.HexColor("#2E6F9E"), borderWidth=1.2,
            borderPadding=5, backColor=colors.HexColor("#EAF2F8"),
            spaceBefore=2, spaceAfter=4, wordWrap="CJK",
            textColor=colors.HexColor("#173A5E"),
        ),
        "caption": ParagraphStyle(
            "KoreanCaption", fontName="MalgunGothic", fontSize=6.3,
            leading=7.8, alignment=TA_CENTER, spaceBefore=1, spaceAfter=3,
            wordWrap="CJK", textColor=colors.HexColor("#697586"),
        ),
        "list": ParagraphStyle(
            "KoreanList", fontName="MalgunGothic", fontSize=7.15,
            leading=8.8, leftIndent=12, firstLineIndent=-8, spaceAfter=1.2,
            wordWrap="CJK", textColor=colors.HexColor("#263646"),
        ),
        "ref": ParagraphStyle(
            "KoreanRef", fontName="MalgunGothic", fontSize=6.9,
            leading=8.5, leftIndent=12, firstLineIndent=-12, spaceAfter=1.5,
            wordWrap="CJK", textColor=colors.HexColor("#263646"),
        ),
        "cell": ParagraphStyle(
            "KoreanCell", fontName="MalgunGothic", fontSize=5.55,
            leading=6.8, wordWrap="CJK", textColor=colors.HexColor("#263646"),
        ),
        "cellhead": ParagraphStyle(
            "KoreanCellHead", fontName="MalgunGothic-Bold", fontSize=5.65,
            leading=6.9, wordWrap="CJK", textColor=colors.white,
        ),
    }


def pdf_image_flowable(path: Path, max_width_cm: float, max_height_cm: float) -> RLImage:
    with PILImage.open(path) as image:
        width_px, height_px = image.size
    ratio = height_px / width_px
    width = max_width_cm * cm
    height = width * ratio
    if height > max_height_cm * cm:
        height = max_height_cm * cm
        width = height / ratio
    return RLImage(str(path), width=width, height=height, hAlign="CENTER")


def add_pdf_table(story, rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    data = []
    for r_idx, row in enumerate(rows):
        style = styles["cellhead"] if r_idx == 0 else styles["cell"]
        data.append([
            RLParagraph(markdown_inline_to_reportlab(row[c] if c < len(row) else ""), style)
            for c in range(ncols)
        ])
    usable = 18.0 * cm
    weights = [1.0] * ncols
    if ncols == 2:
        weights = [1.15, 4.85]
    elif ncols == 4:
        weights = [1.15, 1.75, 2.05, 1.85]
    elif ncols == 5:
        weights = [1.15, 1.65, 1.85, 1.75, 1.55]
    total = sum(weights)
    col_widths = [usable * w / total for w in weights]
    table = RLTable(data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#173A5E")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB5C0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2.8),
        ("TOPPADDING", (0, 0), (-1, -1), 2.0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.0),
    ]
    for r in range(1, len(data)):
        if r % 2 == 0:
            commands.append(("BACKGROUND", (0, r), (-1, r), colors.HexColor("#F2F4F7")))
    table.setStyle(TableStyle(commands))
    story.extend([table, Spacer(1, 2.5)])


def draw_pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    page = canvas.getPageNumber()
    if page > 1:
        canvas.setFont("MalgunGothic", 6.2)
        canvas.setFillColor(colors.HexColor("#697586"))
        canvas.drawString(1.5 * cm, A4[1] - 0.9 * cm, "VALIDITY-GATED INDUSTRIAL DEFECT ACTIVE LEARNING")
    canvas.setFont("MalgunGothic", 6.5)
    canvas.setFillColor(colors.HexColor("#697586"))
    canvas.drawRightString(A4[0] - 1.5 * cm, 0.75 * cm, str(page))
    canvas.restoreState()


def build_pdf(source: Path, output: Path) -> None:
    register_pdf_fonts()
    styles = pdf_styles()
    story = []
    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    seen_title = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            story.append(PageBreak())
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            path = (source.parent / image_match.group(2)).resolve()
            limits = {
                "full_page_validity_gated_architecture.png": (17.8, 9.8),
                "full_page_validity_gated_architecture_v2.png": (17.8, 9.8),
                "deeppcb_prospective_stop_case.png": (16.5, 5.6),
                "validity_gated_algorithm_flowchart.png": (8.2, 11.7),
                "discovery_composition_utility_matrix.png": (16.5, 8.4),
                "paper_framework_temporal_validation.png": (16.5, 8.0),
                "paper_framework_cost_avoidance.png": (16.6, 8.0),
            }
            max_w, max_h = limits.get(path.name, (16.5, 8.5))
            story.append(pdf_image_flowable(path, max_w, max_h))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_pdf_table(story, rows, styles)
            continue
        if line.startswith(">"):
            story.append(RLParagraph(markdown_inline_to_reportlab(line.lstrip("> ")), styles["quote"]))
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not seen_title:
                story.append(RLParagraph(markdown_inline_to_reportlab(text), styles["title"]))
                seen_title = True
            else:
                story.append(RLParagraph(markdown_inline_to_reportlab(text), styles[f"h{level}"]))
            i += 1
            continue
        if seen_title and line.startswith("**Validity-Gated Evaluation"):
            story.append(RLParagraph(markdown_inline_to_reportlab(line.strip("*")), styles["subtitle"]))
            i += 1
            continue
        if seen_title and line.startswith("연구 방향 승인용"):
            story.append(RLParagraph(markdown_inline_to_reportlab(line), styles["meta"]))
            i += 1
            continue
        if line.startswith("**그림 "):
            story.append(RLParagraph(markdown_inline_to_reportlab(line), styles["caption"]))
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            story.append(RLParagraph(markdown_inline_to_reportlab(line), styles["list"]))
            i += 1
            continue
        if line.startswith("- "):
            story.append(RLParagraph(markdown_inline_to_reportlab("• " + line[2:]), styles["list"]))
            i += 1
            continue
        style = styles["ref"] if re.match(r"^\[\d+\]", line) else styles["body"]
        story.append(RLParagraph(markdown_inline_to_reportlab(line), style))
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(
        str(output), pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.35 * cm, bottomMargin=1.25 * cm,
        title="산업 결함 능동학습 후보 신호의 단계적 타당성 평가",
        author="Defect VLM Project",
    )
    pdf.build(story, onFirstPage=draw_pdf_header_footer, onLaterPages=draw_pdf_header_footer)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def write_registry(run_dir: Path, source: Path, output: Path, pdf_output: Path) -> None:
    run_dir.mkdir(parents=True, exist_ok=True)
    sources = [
        (source, "canonical mini-paper source"),
        (ROOT / "docs/research_evolution_and_evidence_freeze_v3_20260718.md", "Evidence Freeze v3 decision"),
        (ROOT / "docs/thesis_claim_boundary_v3_20260718.md", "Evidence Freeze v3 claim boundary"),
        (ROOT / "runs/evidence_freeze_v3_20260718/research_evidence_ledger.csv", "Evidence Freeze v3 ledger including DeepPCB E035-E046"),
        (ROOT / "docs/deeppcb_branch_closure_20260718.md", "DeepPCB exact branch closure"),
        (ROOT / "docs/advisor_decision_brief_deeppcb_closure_20260718.md", "DeepPCB advisor decision brief"),
        (ROOT / "docs/thesis_defense_readiness_v3_20260718.md", "thesis defense non-regression contract"),
        (ROOT / "docs/research_evolution_and_evidence_freeze_v2_20260718.md", "evidence freeze decision"),
        (ROOT / "docs/framework_temporal_validation_decision_20260718.md", "temporal identifiability decision"),
        (ROOT / "docs/framework_branch_timeline_20260718.csv", "branch chronology"),
        (ROOT / "docs/framework_holdout_confusion_matrix_20260718.csv", "holdout identifiability"),
        (ROOT / "docs/framework_cost_avoidance_summary_20260718.csv", "cost lower bound"),
        (ROOT / "docs/thesis_claim_boundary_20260718.md", "claim boundary"),
        (ROOT / "runs/evidence_freeze_v2_20260718/research_evidence_ledger.csv", "frozen evidence ledger"),
        (ROOT / "docs/validity_gated_workflow_algorithm_20260718.md", "algorithm specification"),
        (ROOT / "docs/mini_paper_core_tables_20260718.md", "core tables"),
        (ROOT / "docs/mini_paper_self_persuasion_audit_20260718.md", "self-persuasion audit"),
        (ROOT / "docs/advisor_mini_paper_decision_package_20260718.md", "advisor decision package"),
        (ROOT / "docs/figures/full_page_validity_gated_architecture.png", "architecture figure"),
        (ROOT / "docs/figures/full_page_validity_gated_architecture_v2.png", "architecture figure with DeepPCB prospective STOP case"),
        (ROOT / "docs/figures/deeppcb_prospective_stop_case.png", "DeepPCB prospective authorization STOP case"),
        (ROOT / "docs/figures/validity_gated_algorithm_flowchart.png", "algorithm flowchart"),
        (ROOT / "docs/figures/paper_framework_temporal_validation.png", "paper temporal-identifiability figure"),
        (ROOT / "docs/figures/paper_framework_cost_avoidance.png", "paper cost-containment figure"),
        (ROOT / "docs/figures/paper_framework_advance_stop_timeline.png", "paper branch chronology figure"),
        (output, "generated DOCX"),
        (pdf_output, "generated PDF"),
    ]
    registry = run_dir / "source_registry.csv"
    with registry.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=["path", "role", "exists", "bytes", "sha256"])
        writer.writeheader()
        for path, role in sources:
            exists = path.exists()
            writer.writerow({
                "path": str(path.relative_to(ROOT)) if path.is_relative_to(ROOT) else str(path),
                "role": role,
                "exists": exists,
                "bytes": path.stat().st_size if exists else 0,
                "sha256": sha256(path) if exists else "",
            })

    missing = [
        ("prospective_generic_policy", "Generic gate policy frozen before all evaluated branches", "NOT_AVAILABLE", "predictive screening metrics remain NA"),
        ("stopped_branch_counterfactuals", "Downstream outcomes for branches stopped before training", "STRUCTURALLY_MISSING", "false-stop sensitivity not identifiable"),
        ("robust_positive_selector", "Known-good selector under identical prospective protocol", "NOT_AVAILABLE", "good-selector sensitivity not identifiable"),
        ("independent_production_pools", "Target-blind lot/time/source pool realizations", "NOT_AVAILABLE", "no production generalization claim"),
        ("actual_annotation_cost", "Time and monetary annotation measurements", "NOT_MEASURED", "report model-run lower bound only"),
        ("human_trust", "Inspector trust and explanation usefulness study", "NOT_MEASURED", "no human acceptance claim"),
        ("fn_external_confirmation", "Untouched external FN-repair confirmation", "NOT_AVAILABLE", "FN enrichment remains exploratory"),
    ]
    with (run_dir / "missing_evidence.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["evidence", "required_observation", "status", "claim_effect"])
        writer.writerows(missing)

    config = {
        "package": "mini_paper_validity_gated_industrial_al_v2_20260718",
        "source": str(source.relative_to(ROOT)),
        "docx": str(output.relative_to(ROOT)),
        "pdf": str(pdf_output.relative_to(ROOT)),
        "framework_identity": "retrospective validity-gated empirical evaluation and cost-containment workflow",
        "predictive_policy_claim": False,
        "training_performed": False,
        "inference_performed": False,
        "vlm_calls_performed": False,
        "embedding_extraction_performed": False,
        "selector_implementation_performed": False,
        "fn_screen_performed": False,
        "final_test_used": False,
    }
    (run_dir / "package_config.json").write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")
    commands = [
        ".\\.venv\\Scripts\\python.exe scripts\\04_dcal_xai\\build_mini_paper_figures.py",
        "<bundled-python> scripts\\04_dcal_xai\\build_mini_paper_package.py",
        "<bundled-python> <documents-skill>\\render_docx.py docs\\mini_paper_validity_gated_industrial_al_20260718.docx --output_dir runs\\mini_paper_package_20260718\\docx_render_qa --emit_pdf",
        "<bundled-python> scripts\\04_dcal_xai\\test_mini_paper_package.py",
    ]
    (run_dir / "executed_commands.txt").write_text("\n".join(commands) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--pdf-output", type=Path, default=DEFAULT_PDF)
    parser.add_argument("--run-dir", type=Path, default=DEFAULT_RUN_DIR)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    source = args.source.resolve()
    output = args.output.resolve()
    pdf_output = args.pdf_output.resolve()
    run_dir = args.run_dir.resolve()
    if not source.exists():
        raise FileNotFoundError(source)
    required_images = [
        ROOT / "docs/figures/full_page_validity_gated_architecture_v2.png",
        ROOT / "docs/figures/deeppcb_prospective_stop_case.png",
        ROOT / "docs/figures/validity_gated_algorithm_flowchart.png",
        ROOT / "docs/figures/discovery_composition_utility_matrix.png",
        ROOT / "docs/figures/paper_framework_temporal_validation.png",
        ROOT / "docs/figures/paper_framework_cost_avoidance.png",
    ]
    missing = [str(p) for p in required_images if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing paper figures: " + ", ".join(missing))
    if args.dry_run:
        print(json.dumps({
            "status": "DRY_RUN_PASS",
            "source": str(source),
            "output": str(output),
            "pdf_output": str(pdf_output),
            "required_images": len(required_images),
            "training_performed": False,
            "inference_performed": False,
            "final_test_used": False,
        }, ensure_ascii=False, indent=2))
        return
    build_docx(source, output)
    build_pdf(source, pdf_output)
    write_registry(run_dir, source, output, pdf_output)
    print(json.dumps({
        "status": "DONE",
        "docx": str(output),
        "pdf": str(pdf_output),
        "run_dir": str(run_dir),
        "training_performed": False,
        "inference_performed": False,
        "final_test_used": False,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
